"""Generate NAJIK software development agreement (black-and-white .docx)."""
from datetime import date

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Emu, Pt, RGBColor

OUT = r"D:\NAJIK_Software_Development_Agreement_NepTech_Kalash.docx"
OUT_ALT = r"D:\NAJIK_Contract_NepTech_Kalash_Letterhead.docx"
LOGO_SRC = r"D:\IPO_BULK_APPLY_PROJECT\scripts\neptech_logo.png"
LOGO = r"D:\IPO_BULK_APPLY_PROJECT\scripts\neptech_logo_hd.png"

BLACK = RGBColor(0, 0, 0)
DARK = RGBColor(0x1A, 0x1A, 0x1A)
GRAY = RGBColor(0x3A, 0x3A, 0x3A)
NAVY = RGBColor(0x1A, 0x3A, 0x6B)
EMAIL = "info@neptech.online"


def set_run(run, *, size=11, bold=False, italic=False, color=BLACK, font="Times New Roman"):
    run.bold = bold
    run.italic = italic
    run.font.name = font
    run.font.size = Pt(size)
    run.font.color.rgb = color
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.get_or_add_rFonts()
    rFonts.set(qn("w:ascii"), font)
    rFonts.set(qn("w:hAnsi"), font)
    rFonts.set(qn("w:eastAsia"), font)
    rFonts.set(qn("w:cs"), font)


def p(doc, text="", *, size=11, bold=False, italic=False, center=False, space_after=6, space_before=0, justify=True):
    para = doc.add_paragraph()
    para.paragraph_format.space_after = Pt(space_after)
    para.paragraph_format.space_before = Pt(space_before)
    para.paragraph_format.line_spacing = 1.15
    if center:
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    elif justify:
        para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    if text:
        run = para.add_run(text)
        set_run(run, size=size, bold=bold, italic=italic)
    return para


def add_runs(para, parts):
    """parts: list of (text, bold, italic, size)"""
    for text, bold, italic, size in parts:
        run = para.add_run(text)
        set_run(run, size=size, bold=bold, italic=italic)


def heading(doc, text, number=None):
    label = f"{number}. {text}" if number else text
    return p(doc, label, size=12, bold=True, justify=False, space_before=12, space_after=6)


def bullet(doc, text):
    para = doc.add_paragraph(style="List Bullet")
    para.paragraph_format.space_after = Pt(3)
    para.paragraph_format.left_indent = Cm(0.75)
    para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    run = para.add_run(text)
    set_run(run, size=11)


def numbered(doc, text):
    para = doc.add_paragraph(style="List Number")
    para.paragraph_format.space_after = Pt(3)
    para.paragraph_format.left_indent = Cm(0.75)
    para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    run = para.add_run(text)
    set_run(run, size=11)


def no_table_borders(table):
    tbl = table._tbl
    tblPr = tbl.tblPr
    if tblPr is None:
        tblPr = OxmlElement("w:tblPr")
        tbl.insert(0, tblPr)
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), "nil")
        el.set(qn("w:sz"), "0")
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), "auto")
        borders.append(el)
    tblPr.append(borders)


def set_cell_width(cell, cm):
    cell.width = Cm(cm)
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcW = OxmlElement("w:tcW")
    tcW.set(qn("w:w"), str(int(cm * 567)))
    tcW.set(qn("w:type"), "dxa")
    tcPr.append(tcW)


def para_bottom_line(para, color="2C2C2C", sz="12"):
    pPr = para._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), sz)
    bottom.set(qn("w:space"), "4")
    bottom.set(qn("w:color"), color)
    pBdr.append(bottom)
    pPr.append(pBdr)


def para_top_line(para, color="B0B0B0", sz="6"):
    pPr = para._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    top = OxmlElement("w:top")
    top.set(qn("w:val"), "single")
    top.set(qn("w:sz"), sz)
    top.set(qn("w:space"), "6")
    top.set(qn("w:color"), color)
    pBdr.append(top)
    pPr.append(pBdr)


def add_page_field(paragraph, size=8, color=GRAY, font="Calibri"):
    run = paragraph.add_run()
    set_run(run, size=size, color=color, font=font)
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_begin)
    run._r.append(instr)
    run._r.append(fld_end)


def prepare_logo() -> str:
    """Crop empty padding and upscale so Word does not blur a small PNG."""
    from PIL import Image

    im = Image.open(LOGO_SRC).convert("RGBA")
    bbox = im.getbbox()
    if bbox:
        im = im.crop(bbox)
    # ~300 dpi at 3.6 cm width
    target_w = 900
    scale = max(3, target_w / max(im.width, 1))
    new_size = (int(im.width * scale), int(im.height * scale))
    im = im.resize(new_size, Image.Resampling.LANCZOS)
    im.save(LOGO, "PNG", optimize=False)
    return LOGO


def disable_image_compression(doc: Document) -> None:
    settings = doc.settings.element
    tag = qn("w:doNotAutoCompressPictures")
    if settings.find(tag) is None:
        el = OxmlElement("w:doNotAutoCompressPictures")
        settings.append(el)


def apply_letterhead(section):
    section.top_margin = Cm(4.8)
    section.bottom_margin = Cm(2.8)
    section.left_margin = Cm(2.0)
    section.right_margin = Cm(2.0)
    section.header_distance = Cm(0.5)
    section.footer_distance = Cm(0.5)
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)

    header = section.header
    header.is_linked_to_previous = False
    hp = header.paragraphs[0]
    hp.text = ""
    hp.paragraph_format.space_after = Pt(0)
    hp.paragraph_format.space_before = Pt(0)

    table = header.add_table(rows=1, cols=2, width=Cm(17.0))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    no_table_borders(table)
    left, right = table.rows[0].cells
    set_cell_width(left, 5.2)
    set_cell_width(right, 11.8)

    lp = left.paragraphs[0]
    lp.alignment = WD_ALIGN_PARAGRAPH.LEFT
    lp.paragraph_format.space_after = Pt(0)
    pic = lp.add_run()
    pic.add_picture(LOGO, width=Cm(3.5))

    rp = right.paragraphs[0]
    rp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    rp.paragraph_format.space_before = Pt(8)
    rp.paragraph_format.space_after = Pt(0)
    r1 = rp.add_run("NEPTECH SOLUTIONS")
    set_run(r1, size=16, bold=True, color=DARK, font="Calibri")
    rp2 = right.add_paragraph()
    rp2.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    rp2.paragraph_format.space_before = Pt(2)
    rp2.paragraph_format.space_after = Pt(0)
    r2 = rp2.add_run("Building Digital Solutions")
    set_run(r2, size=11, italic=True, color=GRAY, font="Calibri")

    contact = header.add_paragraph()
    contact.alignment = WD_ALIGN_PARAGRAPH.CENTER
    contact.paragraph_format.space_before = Pt(10)
    contact.paragraph_format.space_after = Pt(4)
    cr = contact.add_run(
        "New Baneshwor, Kathmandu, Nepal  |  +977 9767498797  |  "
        f"{EMAIL}  |  www.neptech.online"
    )
    set_run(cr, size=9, bold=True, color=DARK, font="Calibri")
    para_bottom_line(contact, color="1A1A1A", sz="18")

    footer = section.footer
    footer.is_linked_to_previous = False
    fp0 = footer.paragraphs[0]
    fp0.text = ""
    fp0.paragraph_format.space_before = Pt(0)
    fp0.paragraph_format.space_after = Pt(0)

    ft = footer.add_table(rows=1, cols=3, width=Cm(17.0))
    ft.alignment = WD_TABLE_ALIGNMENT.CENTER
    no_table_borders(ft)
    c0, c1, c2 = ft.rows[0].cells
    set_cell_width(c0, 6.0)
    set_cell_width(c1, 5.5)
    set_cell_width(c2, 5.5)

    def foot_cell(cell, text, align):
        cell.paragraphs[0].text = ""
        para = cell.paragraphs[0]
        para.alignment = align
        para.paragraph_format.space_before = Pt(2)
        para.paragraph_format.space_after = Pt(0)
        para_top_line(para)
        run = para.add_run(text)
        set_run(run, size=9, bold=True, color=DARK, font="Calibri")

    foot_cell(c0, "NepTech Solutions ~ Confidential", WD_ALIGN_PARAGRAPH.LEFT)
    foot_cell(c1, EMAIL, WD_ALIGN_PARAGRAPH.CENTER)
    foot_cell(c2, "9767498797  |  neptech.online", WD_ALIGN_PARAGRAPH.RIGHT)

    pg = footer.add_paragraph()
    pg.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pg.paragraph_format.space_before = Pt(4)
    pg.paragraph_format.space_after = Pt(0)
    pr = pg.add_run("Software Development Agreement  ·  Page ")
    set_run(pr, size=9, bold=True, color=DARK, font="Calibri")
    add_page_field(pg, size=9, color=DARK, font="Calibri")


def hline(para):
    para_bottom_line(para, color="000000", sz="12")


def main():
    prepare_logo()
    doc = Document()
    disable_image_compression(doc)
    for section in doc.sections:
        apply_letterhead(section)

    title = p(doc, "SOFTWARE DEVELOPMENT AGREEMENT", size=16, bold=True, center=True, space_after=4, justify=False)
    hline(title)
    p(
        doc,
        "Project: NAJIK — Local Marketplace Mobile Application and Desktop Admin Panel",
        size=11,
        italic=True,
        center=True,
        space_after=2,
        justify=False,
    )
    p(
        doc,
        f"Agreement Date: {date.today().strftime('%d %B %Y')}     |     Governing Law: Nepal",
        size=10,
        center=True,
        space_after=14,
        justify=False,
    )

    p(
        doc,
        "This Software Development Agreement (“Agreement”) is entered into as of the Agreement Date stated above by and between the parties identified below (each a “Party” and together the “Parties”).",
        space_after=10,
    )

    heading(doc, "PARTIES", 1)
    p(
        doc,
        "1.1  Service Provider. NepTech Solutions, a technology services firm organized and operating in Nepal, with notice email info@neptech.online (hereinafter “Developer” / “NepTech”).",
        space_after=6,
    )
    p(
        doc,
        "1.2  Client. Kalash Financial Solution Pvt. Ltd., a private limited company incorporated in Nepal (hereinafter “Client” / “Kalash”).",
        space_after=6,
    )
    p(
        doc,
        "1.3  Notices. Formal notices under this Agreement shall be sent to the addresses, emails, or registered offices notified in writing by each Party. Until updated, notices may be delivered in person, by email, or by courier with written acknowledgement.",
        space_after=6,
    )

    heading(doc, "RECITALS", 2)
    p(
        doc,
        "The Client wishes to engage the Developer to design, develop, and deliver a Nepal-focused local marketplace product known as “NAJIK” (tagline: Everything Near You), comprising a user-facing mobile application and a desktop-controlled administrative panel. The Developer has the skill to perform such work. The Parties agree to the terms below.",
        space_after=6,
    )

    heading(doc, "DEFINITIONS", 3)
    p(doc, "In this Agreement, unless the context requires otherwise:", space_after=4)
    bullet(doc, "“Software” means the NAJIK user application, desktop admin panel, and related backend application programming interface (API) developed under this Agreement, within the Agreed Scope.")
    bullet(doc, "“Agreed Scope” means the deliverables listed in Clause 5 only, and excludes items listed in Clause 6.")
    bullet(doc, "“Admin Panel” means a web-based control system used from a desktop or laptop computer browser, for the Client’s owners and authorised staff.")
    bullet(doc, "“User App” means the mobile application used by the public to post listings, search, and communicate with other users.")
    bullet(doc, "“Confidential Information” means non-public business, technical, and user data of either Party.")
    bullet(doc, "“Change Request” means any feature, screen, integration, or category not listed in Clause 5.")

    heading(doc, "NATURE OF THE PRODUCT", 4)
    p(
        doc,
        "4.1  NAJIK is a local classifieds / marketplace platform (similar in concept to a Nepal buy–sell–rent–hire notice board). Members of the public may post listings (for example rooms, vehicles, used items). Other users may view listings and contact the poster by in-app chat (and call where enabled). Selected users may display a verified badge after staff review. The Client’s owners and staff manage users, listings, and reports from the Admin Panel on desktop.",
        space_after=6,
    )
    p(
        doc,
        "4.2  The Software is not a full e-commerce checkout store (no shopping cart, warehouse fulfilment, or guaranteed in-app payment for each item sale unless later agreed in a Change Request). Transactions between buyer and seller may occur outside the Software.",
        space_after=6,
    )
    p(
        doc,
        "4.3  The Admin Panel is intended solely for the Client (app owners) and the Client’s employed or contracted staff. Ordinary buyers and sellers shall not receive Admin Panel access.",
        space_after=6,
    )

    heading(doc, "AGREED SCOPE OF WORK (INCLUDED)", 5)
    p(
        doc,
        "For the Contract Price in Clause 8, the Developer shall deliver the following first production version (“Phase 1”). The Client’s longer product documentation may describe a future vision; only the items in this Clause 5 are included in this price.",
        space_after=6,
    )
    p(doc, "5.1  User Mobile Application (Android first)", bold=True, justify=False, space_after=4)
    bullet(doc, "Registration / login (phone or email as technically agreed).")
    bullet(doc, "Home: search, categories, listing cards, location display (manual location if GPS is denied).")
    bullet(doc, "Explore / search with basic filters (category, location, price).")
    bullet(doc, "Post listing wizard: category, photos, title/description, price (NPR), location, contact preference, preview and publish.")
    bullet(doc, "Listing detail: images, price, seller, save, report.")
    bullet(doc, "One-to-one text chat linked to a listing.")
    bullet(doc, "Saved listings and basic profile (photo, name, my listings, logout).")
    bullet(doc, "Core listing types in Phase 1: Used Items, Property (rooms/rent or sale as designed), and Vehicles. Additional categories may be added as Change Requests.")
    p(doc, "5.2  Desktop Admin Panel (mandatory)", bold=True, justify=False, space_after=4, space_before=6)
    bullet(doc, "Secure login for Client staff; access from desktop/laptop web browser.")
    bullet(doc, "Dashboard with basic counts (users, listings, pending/approved). Detailed revenue analytics and multi-role staff matrix beyond Admin / Super Admin may be Phase 2.")
    bullet(doc, "User list: view, suspend / activate.")
    bullet(doc, "Listing moderation: approve, reject (with reason), hide / unhide.")
    bullet(doc, "Manual “Verified” flag for a user or listing (document KYC workflow is Phase 2 unless agreed).")
    bullet(doc, "Simple reports inbox (user reports on listings).")
    bullet(doc, "Two staff roles minimum: Super Admin (Client owner) and Admin/Moderator (staff).")
    p(doc, "5.3  Backend and delivery", bold=True, justify=False, space_after=4, space_before=6)
    bullet(doc, "Hosted API, database, and image storage wired for the above features.")
    bullet(doc, "HTTPS where the Client has provided a domain and server (or Developer-assisted setup billed as reimbursable cost).")
    bullet(doc, "Source code handover after full payment.")
    bullet(doc, "Handover notes: how to log in, publish a listing, and use the Admin Panel.")
    bullet(doc, "Up to two (2) rounds of reasonable revisions on delivered screens within the Agreed Scope, within fourteen (14) days of each milestone demonstration.")

    heading(doc, "EXCLUSIONS (NOT INCLUDED IN NPR 40,000)", 6)
    p(doc, "Unless a written Change Request and extra fee are agreed, the following are out of scope:", space_after=4)
    bullet(doc, "Apple iOS App Store build and Apple Developer enrolment.")
    bullet(doc, "Interactive map search, Google Maps / Mapbox paid quota, and “search this area”.")
    bullet(doc, "Voice messages, in-app wallet, shopping cart, delivery logistics.")
    bullet(doc, "eSewa / Khalti / other payment gateways for listing boosts, subscriptions, or ads.")
    bullet(doc, "Full KYC document centre, six named admin roles, CMS (FAQ/terms editor), advertising campaigns, and advanced charts (DAU, revenue, city performance).")
    bullet(doc, "Jobs marketplace, services marketplace, and business directory as full modules (beyond a simple extra category if capacity allows by written agreement).")
    bullet(doc, "Complete Nepali language translation of every screen.")
    bullet(doc, "Play Store / App Store developer account fees, SMS OTP credits, map API bills, domain, VPS, email, and object storage invoices — these are Client costs (Clause 9).")
    bullet(doc, "Ongoing 24/7 support, new features after Final Acceptance, and content writing (listing photos, legal policies).")

    heading(doc, "CLIENT RESPONSIBILITIES", 7)
    bullet(doc, "Provide branding assets (logo, colours if different from the NAJIK reference), app name confirmation, and sample content within seven (7) days of the Agreement Date.")
    bullet(doc, "Nominate one (1) authorised contact for feedback. Conflicting instructions from multiple staff may delay the timetable.")
    bullet(doc, "Review demonstrations within five (5) working days. Silence may be treated as approval of that milestone.")
    bullet(doc, "Pay invoices on time. Work may pause if any instalment is more than seven (7) days overdue.")
    bullet(doc, "Obtain and pay for domain, hosting, SMS, stores, and third-party APIs, or reimburse the Developer if the Developer pays on the Client’s behalf (with prior written consent).")
    bullet(doc, "Ensure Admin Panel users are only the Client’s owners and staff; the Client is responsible for staff passwords and misuse.")
    bullet(doc, "Supply legally required texts: Terms, Privacy Policy, and marketplace rules (Developer may insert placeholders until the Client provides final legal copy).")

    heading(doc, "CONTRACT PRICE AND PAYMENT", 8)
    p(
        doc,
        "8.1  Contract Price. The all-inclusive professional fee for the Agreed Scope is Nepalese Rupees Forty Thousand only (NPR 40,000.00), exclusive of third-party and government fees in Clause 9.",
        space_after=6,
    )
    p(
        doc,
        "8.2  Advance. Thirty percent (30%) of the Contract Price, being NPR 12,000.00 (Rupees Twelve Thousand only), is payable as a non-refundable mobilisation advance upon signing of this Agreement (except as Clause 14 may require on Developer default). Work shall commence only after the advance is received in the Developer’s nominated account.",
        space_after=6,
    )
    p(
        doc,
        "8.3  Balance. The remaining seventy percent (70%), being NPR 28,000.00 (Rupees Twenty-Eight Thousand only), is payable upon delivery of the Agreed Scope to a staging or production environment and before final source-code handover and Play Store upload assistance (if any). The Developer may withhold production credentials and source code until this balance is paid in full.",
        space_after=6,
    )
    p(
        doc,
        "8.4  Method. Payment shall be made by bank transfer or other lawful channel to NepTech Solutions. A simple receipt or invoice shall be issued. The Client shall bear its own bank charges.",
        space_after=6,
    )
    p(
        doc,
        "8.5  Change Requests. Work outside Clause 5 shall be quoted in writing (time and fee) and is not covered by NPR 40,000. No extra work is obligatory until both Parties sign or email-confirm the quote.",
        space_after=6,
    )
    p(
        doc,
        "8.6  Taxes. If VAT or withholding tax applies under Nepal law, it shall be handled as required by law and shown on the invoice; the Contract Price is the Developer’s professional fee unless the Parties agree otherwise in writing.",
        space_after=6,
    )

    heading(doc, "THIRD-PARTY AND OPERATING COSTS", 9)
    p(
        doc,
        "The Client shall pay, or reimburse against bills: domain registration; VPS/hosting; SSL if not free; database hosting if separate; image storage; transactional email; SMS OTP; Google Play one-time fee; Apple Developer fee (if iOS is later ordered); map API usage; payment gateway charges; and similar. These amounts are not part of NPR 40,000.",
        space_after=6,
    )

    heading(doc, "TIMETABLE", 10)
    p(
        doc,
        "10.1  Target. Subject to timely advance payment and Client feedback, Phase 1 is targeted within approximately six to eight (6–8) weeks from the date the advance is received. This is an estimate, not a penalty-backed guarantee, because listings, chat, and admin moderation depend on Client content and third-party accounts.",
        space_after=6,
    )
    p(
        doc,
        "10.2  Delay by Client. If branding, feedback, or payments are late, the timetable extends day-for-day. If the Client is silent for more than fifteen (15) consecutive days, the Developer may pause the project; a restart fee may apply if pause exceeds thirty (30) days.",
        space_after=6,
    )
    p(
        doc,
        "10.3  Force majeure. Neither Party is liable for delay caused by events beyond reasonable control (including prolonged internet or cloud outage, law change, or natural disaster), provided the affected Party gives prompt notice.",
        space_after=6,
    )

    heading(doc, "ACCEPTANCE", 11)
    p(
        doc,
        "The Client shall have seven (7) days after notice of delivery to report defects that prevent the Agreed Scope from working in a material way. Cosmetic preferences after two revision rounds are not defects. If no written defect list is sent within seven days, the Software is deemed accepted (“Final Acceptance”). Minor bugs after acceptance may be fixed during a fourteen (14) day courtesy warranty for Phase 1 items only.",
        space_after=6,
    )

    heading(doc, "INTELLECTUAL PROPERTY", 12)
    p(
        doc,
        "12.1  Upon full payment of the Contract Price, the Client receives a perpetual licence to use, operate, and modify the custom source code written uniquely for NAJIK under this Agreement, for the Client’s business.",
        space_after=6,
    )
    p(
        doc,
        "12.2  Until full payment, all custom code, designs, and staging systems remain the property of NepTech Solutions. Pre-existing libraries, open-source components, and Developer tools stay subject to their own licences; the Client receives only the rights those licences allow.",
        space_after=6,
    )
    p(
        doc,
        "12.3  The NAJIK name, logo, and brand belong to the Client (or the Client’s licensor). The Developer may show non-confidential screenshots in a portfolio unless the Client objects in writing.",
        space_after=6,
    )

    heading(doc, "CONFIDENTIALITY AND DATA", 13)
    p(
        doc,
        "Each Party shall keep the other’s Confidential Information secret, using at least reasonable care, except where disclosure is required by law. The Client is the data controller for end-user personal data. The Developer shall not sell user data. After the project, the Developer may keep copies only as needed for backups, tax, or legal defence, then delete on reasonable written request where the law allows.",
        space_after=6,
    )

    heading(doc, "WARRANTY AND LIABILITY", 14)
    p(
        doc,
        "14.1  The Developer warrants that Phase 1 will substantially perform the Agreed Scope when used as instructed on supported Android versions and a modern desktop browser (Chrome or Edge) at delivery.",
        space_after=6,
    )
    p(
        doc,
        "14.2  The Software is provided as a marketplace connection tool. The Developer is not responsible for deals between users, fraud by users, unpaid items, or content posted by users. The Client shall moderate listings and users via the Admin Panel.",
        space_after=6,
    )
    p(
        doc,
        "14.3  Except for fraud or wilful misconduct, each Party’s total liability under this Agreement is limited to the Contract Price actually paid. Neither Party is liable for lost profits or indirect loss.",
        space_after=6,
    )

    heading(doc, "TERMINATION", 15)
    p(
        doc,
        "15.1  Either Party may terminate for material breach if the breach is not cured within fifteen (15) days of written notice.",
        space_after=6,
    )
    p(
        doc,
        "15.2  If the Client terminates for convenience after work has started, the advance is retained, and the Client shall pay for work completed beyond the advance at a reasonable time-based rate, capped so that total fees do not exceed NPR 40,000 plus approved Change Requests.",
        space_after=6,
    )
    p(
        doc,
        "15.3  If the Developer abandons the project without cause, the Client may terminate and the Developer shall refund the unused portion of sums paid, after deducting documented work already delivered, as the Parties reasonably agree or as a court determines.",
        space_after=6,
    )

    heading(doc, "NON-SOLICITATION AND RELATIONSHIP", 16)
    p(
        doc,
        "This Agreement does not create a partnership, joint venture, or employment. The Developer is an independent contractor. For twelve (12) months after this Agreement, neither Party shall solicit the other’s employees introduced solely through this project, except by general public advertisement.",
        space_after=6,
    )

    heading(doc, "GENERAL", 17)
    p(
        doc,
        "17.1  Entire agreement. This document is the entire agreement for Phase 1 and supersedes prior oral discussions. Amendments must be in writing (including email) and signed or clearly confirmed by both Parties.",
        space_after=6,
    )
    p(
        doc,
        "17.2  Severability. If any clause is unenforceable, the rest remains in force.",
        space_after=6,
    )
    p(
        doc,
        "17.3  Assignment. Neither Party may assign this Agreement without the other’s written consent, except that the Client may assign to an affiliate that assumes all obligations.",
        space_after=6,
    )
    p(
        doc,
        "17.4  Governing law and disputes. This Agreement is governed by the laws of Nepal. The Parties shall first attempt good-faith negotiation for fifteen (15) days. Thereafter, courts of competent jurisdiction in Nepal shall have exclusive jurisdiction.",
        space_after=6,
    )
    p(
        doc,
        "17.5  Counterparts and digital signature. This Agreement may be signed in two counterparts (including scanned PDF, printed copies, or electronic / digital signature). An electronic or digital signature shall have the same legal effect as a handwritten signature. Together they form one instrument.",
        space_after=6,
    )

    heading(doc, "SIGNATURES", 18)
    p(
        doc,
        "IN WITNESS WHEREOF the Parties have executed this Agreement on the dates written below, intending to be legally bound.",
        space_after=10,
    )
    p(
        doc,
        "THIS CONTRACT IS SIGNED DIGITALLY",
        size=13,
        bold=True,
        center=True,
        justify=False,
        space_before=4,
        space_after=6,
    )
    p(
        doc,
        "The Parties agree that this Agreement is executed by electronic / digital signature. Such signature is valid, binding, and enforceable to the same extent as an original wet-ink signature. No physical (paper) signature is required for this Agreement to take effect, provided each Party’s authorised signatory is identified below and the signed electronic copy is exchanged by email (including info@neptech.online) or other agreed channel.",
        space_after=14,
    )

    # Signature table
    table = doc.add_table(rows=2, cols=2)
    table.autofit = True
    left, right = table.rows[0].cells
    left_b, right_b = table.rows[1].cells

    def fill_sign(cell, title, company, extra_lines):
        cell.text = ""
        t = cell.paragraphs[0]
        t.alignment = WD_ALIGN_PARAGRAPH.LEFT
        r = t.add_run(title)
        set_run(r, size=11, bold=True)
        p2 = cell.add_paragraph()
        r2 = p2.add_run(company)
        set_run(r2, size=11, bold=True)
        for line in extra_lines:
            px = cell.add_paragraph()
            rx = px.add_run(line)
            set_run(rx, size=11)

    fill_sign(
        left,
        "FOR THE DEVELOPER",
        "NepTech Solutions",
        [
            "",
            "Authorised signatory: ________________",
            "",
            "Name: ______________________________",
            "",
            "Designation: ________________________",
            "",
            "Date: ______________________________",
            "",
            "Digital / e-signature: ________________",
            "",
            "Company stamp (if any):",
            "",
            "",
            "________________________________",
        ],
    )
    fill_sign(
        right,
        "FOR THE CLIENT",
        "Kalash Financial Solution Pvt. Ltd.",
        [
            "",
            "Authorised signatory: ________________",
            "",
            "Name: ______________________________",
            "",
            "Designation: ________________________",
            "",
            "Date: ______________________________",
            "",
            "Digital / e-signature: ________________",
            "",
            "Company stamp (if any):",
            "",
            "",
            "________________________________",
        ],
    )

    p(doc, "", space_after=10)
    p(
        doc,
        "Witness 1: Name _____________________  Signature _____________________  Date _____________",
        size=10,
        justify=False,
        space_after=8,
    )
    p(
        doc,
        "Witness 2: Name _____________________  Signature _____________________  Date _____________",
        size=10,
        justify=False,
        space_after=12,
    )
    end = p(
        doc,
        "— End of Agreement —",
        size=10,
        italic=True,
        center=True,
        justify=False,
        space_before=8,
    )

    try:
        doc.save(OUT)
        saved = OUT
    except PermissionError:
        doc.save(OUT_ALT)
        saved = OUT_ALT
    copy_to = r"D:\IPO_BULK_APPLY_PROJECT\NAJIK_Software_Development_Agreement_NepTech_Kalash.docx"
    try:
        import shutil

        shutil.copyfile(saved, copy_to)
    except OSError:
        pass
    print("Wrote", saved)


if __name__ == "__main__":
    main()
