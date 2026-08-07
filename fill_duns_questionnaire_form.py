"""Fill D&B SAMEA questionnaire in official-like form layout."""
from docx import Document
from docx.shared import Pt, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


def set_cell_text(cell, text, bold=False, size=9):
    cell.text = ""
    p = cell.paragraphs[0]
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    run.font.name = "Calibri"


def add_label_value_table(doc, pairs, cols=2):
    """pairs: list of (label, value)"""
    rows = (len(pairs) + cols - 1) // cols
    table = doc.add_table(rows=rows, cols=cols * 2)
    table.style = "Table Grid"
    i = 0
    for r in range(rows):
        for c in range(cols):
            if i >= len(pairs):
                break
            label, value = pairs[i]
            set_cell_text(table.rows[r].cells[c * 2], label, bold=True, size=8)
            set_cell_text(table.rows[r].cells[c * 2 + 1], value, size=9)
            i += 1
    doc.add_paragraph("")


def main():
    doc = Document()
    for section in doc.sections:
        section.top_margin = Cm(1.2)
        section.bottom_margin = Cm(1.2)
        section.left_margin = Cm(1.5)
        section.right_margin = Cm(1.5)

    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(9)

    # Header
    h = doc.add_paragraph()
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = h.add_run("DUN & BRADSTREET SOUTH ASIA MIDDLE EAST LIMITED")
    r.bold = True
    r.font.size = Pt(12)

    h2 = doc.add_paragraph()
    h2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = h2.add_run("QUESTIONNAIRE")
    r2.bold = True
    r2.font.size = Pt(14)

    note = doc.add_paragraph()
    note.add_run(
        "Note: Headings marked ** are mandatory. Handwritten questionnaires are NOT ACCEPTABLE.\n"
        "Purpose: FREE standard D-U-N-S only (Google Play Organization). No paid Certificate/DRS."
    ).font.size = Pt(8)

    # ---- PAGE 1 style content ----
    doc.add_heading("Company Identification", level=2)

    t = doc.add_table(rows=8, cols=4)
    t.style = "Table Grid"
    rows_data = [
        ("Legally Registered Company Name:**", "Kalash Financial Solution Private Limited", "", ""),
        ("Native / Local Name:", "कलश फाइनान्सियल सोलुसन्स प्रा.लि.", "Trading Style:", "NEPSE GHAR"),
        ("Acronym:", "KFS", "", ""),
        (
            "Office Physical Address:**",
            "Ward No. 3, Lahan Municipality, Siraha District, Nepal",
            "Land Mark:",
            "[CONFIRM]",
        ),
        ("City:", "Lahan", "Country:", "Nepal"),
        (
            "Mailing / P.O. BOX Add:",
            "Same as Office Physical Address, Lahan, Nepal",
            "",
            "",
        ),
        (
            "Registered Office Add:**",
            "Ward No. 3, Lahan Municipality, Siraha District, Nepal",
            "City / Country:",
            "Lahan / Nepal",
        ),
        ("Telephone:", "+977 9709133067", "Mobile:", "+977 9709133067"),
    ]
    for i, row in enumerate(rows_data):
        for j, val in enumerate(row):
            set_cell_text(t.rows[i].cells[j], val, bold=(j % 2 == 0), size=8)

    t2 = doc.add_table(rows=1, cols=4)
    t2.style = "Table Grid"
    set_cell_text(t2.rows[0].cells[0], "Fax:", bold=True, size=8)
    set_cell_text(t2.rows[0].cells[1], "N/A", size=9)
    set_cell_text(t2.rows[0].cells[2], "Email Address:", bold=True, size=8)
    set_cell_text(
        t2.rows[0].cells[3],
        "kalash.financialsolutions@gmail.com / info@nepseghar.com",
        size=8,
    )

    t3 = doc.add_table(rows=1, cols=2)
    t3.style = "Table Grid"
    set_cell_text(t3.rows[0].cells[0], "Website:", bold=True, size=8)
    set_cell_text(t3.rows[0].cells[1], "https://www.nepseghar.com", size=9)

    doc.add_heading("Company Details **", level=2)
    cd = doc.add_table(rows=6, cols=4)
    cd.style = "Table Grid"
    cd_rows = [
        ("Year of Establishment:**", "2026 (27 April 2026 AD / 2083-01-14 BS)", "Legal Structure:**", "Pvt. Ltd."),
        ("Employees at Head office only:", "5", "Employees strength total group:", "5"),
        ("Employees strength Yr-1:", "N/A (new company 2026)", "Employees strength Yr-2:", "N/A"),
        ("Total employees HO & Branches:", "5", "Currency:", "NPR"),
        ("Estimated Sales – Current Year:", "[CONFIRM NPR]", "Estimated Sales – Previous Year:", "N/A"),
        (
            "Business Activities (LOB):",
            "Others (X) — Software / IT / Digital share-market tools. NOT Manufacturer/Wholesaler/Retailer. (Template Agent X ignored)",
            "",
            "",
        ),
    ]
    for i, row in enumerate(cd_rows):
        for j, val in enumerate(row):
            set_cell_text(cd.rows[i].cells[j], val, bold=(j % 2 == 0), size=8)

    lob = doc.add_table(rows=4, cols=2)
    lob.style = "Table Grid"
    set_cell_text(lob.rows[0].cells[0], "1) Primary Line of Business:", bold=True, size=8)
    set_cell_text(
        lob.rows[0].cells[1],
        "Computer programming; NEPSE GHAR mobile app — market data, MeroShare/IPO tools, investor utilities",
        size=8,
    )
    set_cell_text(lob.rows[1].cells[0], "2) Secondary Line of Business:", bold=True, size=8)
    set_cell_text(
        lob.rows[1].cells[1],
        "Web portals; data processing/hosting; information services related to capital-market tools",
        size=8,
    )
    set_cell_text(lob.rows[2].cells[0], "Product / Brand Details:", bold=True, size=8)
    set_cell_text(lob.rows[2].cells[1], "Brand: NEPSE GHAR (digital app/tools). No physical products.", size=8)
    set_cell_text(lob.rows[3].cells[0], "Premises used as / Ownership:", bold=True, size=8)
    set_cell_text(
        lob.rows[3].cells[1],
        "Office. Ownership: [CONFIRM – Leased / Owned / Rented / Shared]",
        size=8,
    )

    doc.add_heading("Registration Details **", level=2)
    reg = doc.add_table(rows=5, cols=4)
    reg.style = "Table Grid"
    reg_rows = [
        ("Type of registration:", "Private Limited Company (Companies Act, 2006)", "Trade License / Reg. No:", "390691/82/83"),
        ("Old Registration Number:", "N/A", "Date of Registration:", "27/04/2026"),
        ("Date of Expiry:", "N/A (ongoing company registration)", "Chamber of commerce No:", "[CONFIRM or N/A]"),
        ("PAN Number:", "623600069", "PAN Office:", "IRO Lahan"),
        ("", "", "", ""),
    ]
    for i, row in enumerate(reg_rows):
        for j, val in enumerate(row):
            set_cell_text(reg.rows[i].cells[j], val, bold=(j % 2 == 0 and bool(val)), size=8)

    doc.add_heading("Capital Structure (Currency: NPR)", level=2)
    cap = doc.add_table(rows=5, cols=2)
    cap.style = "Table Grid"
    for i, (a, b) in enumerate(
        [
            ("Authorized Capital – Equity", "[CONFIRM]"),
            ("Issued Capital – Equity", "[CONFIRM]"),
            ("Paid Up Capital – Equity", "[CONFIRM]"),
            ("Partnership / Proprietorship Capital", "N/A (Pvt. Ltd.)"),
            ("", ""),
        ]
    ):
        set_cell_text(cap.rows[i].cells[0], a, bold=True, size=8)
        set_cell_text(cap.rows[i].cells[1], b, size=9)

    doc.add_heading("Management Details (Directors / Key Personnel)", level=2)
    mgmt = doc.add_table(rows=3, cols=4)
    mgmt.style = "Table Grid"
    headers = ["Name", "Designation", "Contact Details & Email", "Nationality"]
    for j, htxt in enumerate(headers):
        set_cell_text(mgmt.rows[0].cells[j], htxt, bold=True, size=8)
    set_cell_text(mgmt.rows[1].cells[0], "Sumit Saphi", size=8)
    set_cell_text(mgmt.rows[1].cells[1], "Director / Authorized Contact [CONFIRM title]", size=8)
    set_cell_text(
        mgmt.rows[1].cells[2],
        "+9779709133067 | kalash.financialsolutions@gmail.com",
        size=7,
    )
    set_cell_text(mgmt.rows[1].cells[3], "Nepali", size=8)
    set_cell_text(mgmt.rows[2].cells[0], "[CONFIRM other directors if any]", size=8)
    set_cell_text(mgmt.rows[2].cells[1], "", size=8)
    set_cell_text(mgmt.rows[2].cells[2], "", size=8)
    set_cell_text(mgmt.rows[2].cells[3], "", size=8)

    doc.add_heading("Bank Details", level=2)
    bank = doc.add_table(rows=2, cols=3)
    bank.style = "Table Grid"
    for j, htxt in enumerate(["Key Bank", "Branch", "Address"]):
        set_cell_text(bank.rows[0].cells[j], htxt, bold=True, size=8)
    set_cell_text(bank.rows[1].cells[0], "[CONFIRM]", size=8)
    set_cell_text(bank.rows[1].cells[1], "[CONFIRM]", size=8)
    set_cell_text(bank.rows[1].cells[2], "[CONFIRM]", size=8)

    doc.add_page_break()

    # ---- PAGE 2 ----
    doc.add_heading("Shareholding Pattern as on (Date: 06 / 08 / 2026)", level=2)
    sh = doc.add_table(rows=3, cols=5)
    sh.style = "Table Grid"
    for j, htxt in enumerate(
        ["Name of the Shareholders", "Local Sponsor (Y/N)", "Address / Location", "Nationality", "% of Total Equity"]
    ):
        set_cell_text(sh.rows[0].cells[j], htxt, bold=True, size=7)
    set_cell_text(sh.rows[1].cells[0], "Sumit Saphi [CONFIRM]", size=8)
    set_cell_text(sh.rows[1].cells[1], "Y", size=8)
    set_cell_text(sh.rows[1].cells[2], "Nepal", size=8)
    set_cell_text(sh.rows[1].cells[3], "Nepali", size=8)
    set_cell_text(sh.rows[1].cells[4], "[CONFIRM %]", size=8)
    set_cell_text(sh.rows[2].cells[0], "[CONFIRM other shareholders]", size=8)

    doc.add_heading("Types Of Customers", level=2)
    doc.add_paragraph(
        "Retail investors / app users; individuals and small businesses using Nepali share-market digital tools "
        "(not Manufacturers/Govt exporters primarily)."
    ).runs[0].font.size = Pt(8)
    cust = doc.add_table(rows=2, cols=2)
    cust.style = "Table Grid"
    set_cell_text(cust.rows[0].cells[0], "Total No of Customers:", bold=True, size=8)
    set_cell_text(cust.rows[0].cells[1], "[CONFIRM]  ☐ 0-10  ☐ 10-25  ☐ 25-50  ☐ More Than 50", size=8)
    set_cell_text(cust.rows[1].cells[0], "Total No of Suppliers:", bold=True, size=8)
    set_cell_text(
        cust.rows[1].cells[1],
        "[CONFIRM]  ☐ 0-10  ☐ 10-25  ☐ 25-50  ☐ More Than 50  (cloud, APIs, app stores)",
        size=8,
    )

    doc.add_heading("Purchases Details", level=2)
    pur = doc.add_table(rows=4, cols=2)
    pur.style = "Table Grid"
    set_cell_text(pur.rows[0].cells[0], "Local %", bold=True, size=8)
    set_cell_text(pur.rows[0].cells[1], "[CONFIRM] %", size=8)
    set_cell_text(pur.rows[1].cells[0], "International %", bold=True, size=8)
    set_cell_text(pur.rows[1].cells[1], "[CONFIRM] %", size=8)
    set_cell_text(pur.rows[2].cells[0], "Local Credit Terms", bold=True, size=8)
    set_cell_text(pur.rows[2].cells[1], "Cash ~100% (digital/service purchases). Credit/LC: N/A or minimal", size=8)
    set_cell_text(pur.rows[3].cells[0], "International Credit Terms", bold=True, size=8)
    set_cell_text(
        pur.rows[3].cells[1],
        "Countries: [CONFIRM if any]. Cash / Wire transfer as applicable. LC: N/A",
        size=8,
    )
    pe = doc.add_table(rows=2, cols=2)
    pe.style = "Table Grid"
    set_cell_text(pe.rows[0].cells[0], "Local Purchase Value (Estimated):", bold=True, size=8)
    set_cell_text(pe.rows[0].cells[1], "[CONFIRM NPR]", size=8)
    set_cell_text(pe.rows[1].cells[0], "International Import Value (Estimated):", bold=True, size=8)
    set_cell_text(pe.rows[1].cells[1], "[CONFIRM NPR / N/A]", size=8)

    doc.add_heading("Sales Details", level=2)
    sales = doc.add_table(rows=4, cols=2)
    sales.style = "Table Grid"
    set_cell_text(sales.rows[0].cells[0], "Local %", bold=True, size=8)
    set_cell_text(sales.rows[0].cells[1], "Approx. 90%+ (Nepal users) [CONFIRM]", size=8)
    set_cell_text(sales.rows[1].cells[0], "International %", bold=True, size=8)
    set_cell_text(sales.rows[1].cells[1], "Approx. 0–10% [CONFIRM]", size=8)
    set_cell_text(sales.rows[2].cells[0], "Local Credit Terms", bold=True, size=8)
    set_cell_text(sales.rows[2].cells[1], "Cash / digital payments ~100%", size=8)
    set_cell_text(sales.rows[3].cells[0], "International Credit Terms", bold=True, size=8)
    set_cell_text(sales.rows[3].cells[1], "Others: WIRE TRANSFER / digital as applicable", size=8)
    se = doc.add_table(rows=2, cols=2)
    se.style = "Table Grid"
    set_cell_text(se.rows[0].cells[0], "Local Sales Value (Estimated):", bold=True, size=8)
    set_cell_text(se.rows[0].cells[1], "[CONFIRM NPR]", size=8)
    set_cell_text(se.rows[1].cells[0], "International Export Value (Estimated):", bold=True, size=8)
    set_cell_text(se.rows[1].cells[1], "N/A or minimal [CONFIRM]", size=8)

    doc.add_heading("Changes in the Company", level=2)
    doc.add_paragraph(
        "None since incorporation (new company, April 2026). "
        "No change in Legal Name / Shareholders / LOB / Legal Structure / Merger / Address / Capital."
    ).runs[0].font.size = Pt(8)

    doc.add_heading("Related Companies (Corporate Linkage) or Branches **", level=2)
    rel = doc.add_table(rows=2, cols=4)
    rel.style = "Table Grid"
    for j, htxt in enumerate(["Name of the Company", "Address", "Nature (HQ/Parent/Subsidiary/Branch)", "% of Shares"]):
        set_cell_text(rel.rows[0].cells[j], htxt, bold=True, size=7)
    set_cell_text(rel.rows[1].cells[0], "None", size=8)
    set_cell_text(rel.rows[1].cells[1], "N/A", size=8)
    set_cell_text(rel.rows[1].cells[2], "N/A — single company, no branches yet [CONFIRM]", size=8)
    set_cell_text(rel.rows[1].cells[3], "N/A", size=8)

    doc.add_page_break()

    # ---- PAGE 3 ----
    doc.add_heading("Number of Facilities", level=2)
    fac = doc.add_table(rows=5, cols=6)
    fac.style = "Table Grid"
    facilities = [
        ("Warehouses", "0", "Computers", "[CONFIRM #]", "Vessels", "0"),
        ("Cars", "0", "Offices", "1", "Machines", "0"),
        ("Vans", "0", "Rooms", "[CONFIRM]", "Classrooms", "0"),
        ("Trucks", "0", "Stores", "0", "", ""),
        ("", "", "", "", "", ""),
    ]
    for i, row in enumerate(facilities):
        for j, val in enumerate(row):
            set_cell_text(fac.rows[i].cells[j], val, bold=(j % 2 == 0 and bool(val)), size=8)

    doc.add_heading("Production Details (in case of Manufacturer)", level=2)
    prod = doc.add_table(rows=2, cols=2)
    prod.style = "Table Grid"
    set_cell_text(prod.rows[0].cells[0], "Type of Goods", bold=True, size=8)
    set_cell_text(prod.rows[0].cells[1], "Installed Facility", bold=True, size=8)
    set_cell_text(prod.rows[1].cells[0], "N/A — not a manufacturer (software/digital services)", size=8)
    set_cell_text(prod.rows[1].cells[1], "N/A", size=8)

    doc.add_heading("Other details", level=2)
    other = doc.add_table(rows=5, cols=2)
    other.style = "Table Grid"
    other_rows = [
        ("Founder Details", "Sumit Saphi"),
        ("Certifications", "N/A currently"),
        ("Lawyer Details", "[CONFIRM or N/A]"),
        ("Auditor Details", "[CONFIRM or N/A]"),
        ("Insurance Details", "[CONFIRM or N/A]"),
    ]
    for i, (a, b) in enumerate(other_rows):
        set_cell_text(other.rows[i].cells[0], a, bold=True, size=8)
        set_cell_text(other.rows[i].cells[1], b, size=8)

    doc.add_heading("Companies Future Plans", level=2)
    doc.add_paragraph(
        "Publish and grow NEPSE GHAR on Google Play Store; expand digital share-market tools "
        "and investor utilities for users in Nepal."
    ).runs[0].font.size = Pt(9)

    doc.add_heading("Interviewer details", level=2)
    iv = doc.add_table(rows=1, cols=6)
    iv.style = "Table Grid"
    set_cell_text(iv.rows[0].cells[0], "Name:", bold=True, size=8)
    set_cell_text(iv.rows[0].cells[1], "Sumit Saphi", size=8)
    set_cell_text(iv.rows[0].cells[2], "Designation:", bold=True, size=8)
    set_cell_text(iv.rows[0].cells[3], "Director / Authorized Representative", size=8)
    set_cell_text(iv.rows[0].cells[4], "Date:", bold=True, size=8)
    set_cell_text(iv.rows[0].cells[5], "06/08/2026", size=8)

    doc.add_heading("Enclosed Details", level=2)
    doc.add_paragraph("☑ Registry / Trade License / Company Registration Certificate (OCR) — 390691/82/83")
    doc.add_paragraph("☑ PAN Registration Certificate — 623600069")
    doc.add_paragraph("☐ Address proof (electricity/water/internet bill or notarized rental) — TO ATTACH")
    doc.add_paragraph("☐ 3 Years Audited Financials — N/A (new company incorporated April 2026)")
    doc.add_paragraph("☐ Site Photographs — optional")
    doc.add_paragraph("☑ Others: Completed questionnaire (this document)")

    foot = doc.add_paragraph()
    foot.add_run(
        "\nREQUEST: Please process STANDARD FREE D-U-N-S Number only (~30 working days). "
        "We do not require paid D-U-N-S Certificate ($750), DRS seal, or expedited packages."
    ).bold = True

    out = r"D:\Kalash_Financial_DUNS_Questionnaire_FILLED_FORM.docx"
    doc.save(out)
    print("SAVED", out)


if __name__ == "__main__":
    main()
