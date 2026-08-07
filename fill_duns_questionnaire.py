from docx import Document
from docx.shared import Pt


def main() -> None:
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    def h(text: str) -> None:
        doc.add_heading(text, level=1)

    def row(label: str, value: str) -> None:
        p = doc.add_paragraph()
        run = p.add_run(f"{label} ")
        run.bold = True
        p.add_run(value)

    doc.add_heading(
        "DUN & BRADSTREET SAMEA — COMPANY INFORMATION QUESTIONNAIRE (FILLED)",
        0,
    )
    p = doc.add_paragraph()
    p.add_run("Note: ").bold = True
    p.add_run(
        "Filled for free standard D-U-N-S registration for Google Play Organization "
        "verification. Fields marked [CONFIRM] need client confirmation. Do NOT pay for "
        "Certificate/DRS unless the client requests it."
    )

    h("1. Company identity")
    row("Legally Registered Company Name:", "Kalash Financial Solution Private Limited")
    row(
        "Native / Local Name:",
        "कलश फाइनान्सियल सोलुसन्स प्रा.लि. (Kalash Financial Solutions Pvt. Ltd.)",
    )
    row("Trading Style:", "NEPSE GHAR")
    row("Acronym:", "KFS / NEPSE GHAR")

    h("2. Addresses")
    row(
        "Office Physical Address:",
        "Ward No. 3, Lahan Municipality, Siraha District, Nepal",
    )
    row("Land Mark:", "[CONFIRM – e.g. near landmark]")
    row("City:", "Lahan")
    row("Country:", "Nepal")
    row("Mailing / P.O. BOX Add:", "Same as office physical address")
    row(
        "Registered Office Add:",
        "Ward No. 3, Lahan Municipality, Siraha District, Nepal",
    )

    h("3. Contacts")
    row("Telephone:", "+977 9709133067")
    row("Mobile:", "+977 9709133067")
    row("Fax:", "N/A")
    row(
        "Email Address:",
        "kalash.financialsolutions@gmail.com / info@nepseghar.com",
    )
    row("Website:", "https://www.nepseghar.com")

    h("4. Company details")
    row(
        "Year of Establishment:",
        "2026 (Incorporated 27 April 2026 AD / 2083-01-14 BS)",
    )
    row("Legal Structure:", "Private Limited (Pvt. Ltd.)")
    row("Employees at Head office only:", "5")
    row(
        "Employees strength previous Year-1:",
        "N/A (company newly incorporated in 2026)",
    )
    row(
        "Employees strength previous Year-2:",
        "N/A (company newly incorporated in 2026)",
    )
    row("Total employees (HO & Branches):", "5")
    row("Employees strength total group:", "5")
    row("Estimated Sales Value – Current Year:", "[CONFIRM – NPR amount]")
    row("Estimated Sales Value – Previous Year:", "N/A (new company)")
    row("Currency:", "NPR (Nepalese Rupee)")

    h("5. Business activities (LOB)")
    row(
        "Business type checkboxes:",
        "Others (Software / IT / Digital financial market tools). Not a manufacturer/wholesaler/retailer.",
    )
    row(
        "1) Primary Line of Business:",
        "Computer programming; development of mobile/web applications for Nepali share market "
        "(NEPSE GHAR) — market data, MeroShare/IPO tools, investor utilities",
    )
    row(
        "2) Secondary Line of Business:",
        "Web portals; data processing/hosting related activities; information services related to capital market tools",
    )
    row(
        "Dealing with Product / Brand Details:",
        "Brand: NEPSE GHAR (Android app / digital tools). No physical product brands.",
    )
    row("Premises used as:", "Office")
    row("Ownership of premises:", "[CONFIRM – Owned / Leased / Rented / Shared]")

    h("6. Registration details")
    row(
        "Type of registration:",
        "Private Limited Company (Companies Act, 2006 – Nepal)",
    )
    row("Trade License / Company Reg. No:", "390691/82/83")
    row("Old Registration Number:", "N/A")
    row(
        "Date of Registration & Expiry:",
        "Registered: 27 April 2026 AD. Expiry: N/A (company registration ongoing)",
    )
    row("Chamber of commerce No:", "[CONFIRM or N/A]")
    row("Authorized Capital:", "[CONFIRM – amount + currency NPR]")
    row("Issued Capital:", "[CONFIRM]")
    row("Paid Up Capital:", "[CONFIRM]")
    row("Equity / Partnership Capital:", "[CONFIRM if applicable]")
    row("PAN Number:", "623600069")
    row("PAN Issuing Office:", "Inland Revenue Office, Lahan")

    h("7. Management details")
    row("Name:", "Sumit Saphi")
    row("Designation:", "Director / Authorized Contact [CONFIRM exact title]")
    row(
        "Contact Details & Email Id:",
        "+977 9709133067 | kalash.financialsolutions@gmail.com",
    )
    row("Nationality:", "Nepali")
    doc.add_paragraph(
        "Additional directors/partners: [CONFIRM – list all directors from company papers with % share]"
    )

    h("8. Bank details")
    row("Key Bank:", "[CONFIRM]")
    row("Branch:", "[CONFIRM]")
    row("Address:", "[CONFIRM]")

    h("9. Shareholding pattern")
    row("As on date:", "06 / 08 / 2026")
    row("Shareholder 1 – Name:", "Sumit Saphi [CONFIRM]")
    row("Local Sponsor (Y/N):", "Y")
    row("Address / Location:", "Nepal")
    row("% of Total:", "[CONFIRM %]")
    doc.add_paragraph("Add other shareholders if any: [CONFIRM]")

    h("10. Customers & suppliers")
    row(
        "Types of Customers:",
        "Retail investors / app users; small businesses using share-market tools",
    )
    row("Total No of Customers:", "[CONFIRM] / currently early stage")
    row("More Than 50:", "[CONFIRM Y/N]")
    row(
        "Total No of Suppliers:",
        "[CONFIRM] (cloud hosting, APIs, app stores, etc.)",
    )
    row("Local purchases %:", "[CONFIRM]")
    row("International purchases %:", "[CONFIRM]")
    row(
        "Credit Terms – Cash / Credit / LC / Others:",
        "Primarily cash / digital payments; Wire transfer for services as applicable",
    )
    row("Local Purchase Value (Estimated):", "[CONFIRM]")
    row("International Import Value (Estimated):", "[CONFIRM]")
    row("Local Sales Value (Estimated):", "[CONFIRM]")
    row("International Export Value (Estimated):", "N/A or minimal [CONFIRM]")

    h("11. Changes in the company")
    row(
        "Change in Legal Name / Shareholders / LOB / Legal Structure / Merger / Address / Capital:",
        "None since incorporation (new company, April 2026)",
    )

    h("12. Related companies / branches")
    row("Related Companies / Branches:", "None [CONFIRM]")

    h("13. Facilities")
    row("Offices:", "1 (Lahan, Siraha)")
    row("Warehouses / Machines / Trucks / Stores:", "N/A (software company)")
    row("Computers:", "[CONFIRM approximate number]")

    h("14. Other details")
    row("Founder Details:", "Sumit Saphi")
    row("Certifications:", "N/A currently")
    row("Lawyer Details:", "[CONFIRM or N/A]")
    row("Auditor Details:", "[CONFIRM or N/A]")
    row("Insurance Details:", "[CONFIRM or N/A]")
    row(
        "Companies Future Plans:",
        "Publish and grow NEPSE GHAR on Google Play; expand digital share-market tools for Nepal investors",
    )

    h("15. Interviewer / applicant details")
    row("Name:", "Sumit Saphi")
    row("Designation:", "Director / Authorized Representative")
    row("Date:", "6 August 2026")

    h("16. Enclosed documents")
    doc.add_paragraph(
        "☑ Company Registration / Incorporation Certificate (OCR) — Reg. No. 390691/82/83"
    )
    doc.add_paragraph("☑ PAN Registration Certificate — PAN 623600069")
    doc.add_paragraph(
        "☐ Address proof (electricity / water / internet bill or notarized rental agreement) — TO ATTACH"
    )
    doc.add_paragraph("☐ 3 Years Audited Financials — N/A (new company)")
    doc.add_paragraph("☐ Site Photographs — optional")

    h("Important note for Dun & Bradstreet")
    doc.add_paragraph(
        "We request STANDARD FREE D-U-N-S Number registration only (approx. 30 working days). "
        "We do not require paid D-U-N-S Certificate, DRS seal, or expedited packages at this time. "
        "Purpose: Google Play Console Organization account verification for publishing the NEPSE GHAR Android application."
    )

    out = r"D:\Kalash_Financial_DUNS_Questionnaire_FILLED.docx"
    doc.save(out)
    print("SAVED", out)


if __name__ == "__main__":
    main()
